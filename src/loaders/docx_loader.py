"""
Word Document Loader for Enterprise AI Knowledge Assistant.
Extracts text, tables and metadata from .docx files.
"""

from docx import Document
from pathlib import Path
from typing import List, Dict, Any
from src.utils.logger import get_loader_logger

logger = get_loader_logger()


class DocxLoader:
    """
    Loads and extracts content from Word (.docx) files.

    Features:
        ✅ Extracts paragraphs and headings
        ✅ Extracts tables as readable text
        ✅ Captures document metadata
        ✅ Handles corrupted files gracefully
    """

    def __init__(self):
        self.supported_extensions = [".docx"]

    def load(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Load a Word file and extract all content.

        Args:
            file_path: Path to the .docx file

        Returns:
            List of documents with text and metadata
        """
        path = Path(file_path)

        # ── Validate file ────────────────────────────────────────
        if not path.exists():
            logger.error(f"Word file not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        if path.suffix.lower() != ".docx":
            logger.error(f"Not a Word file: {file_path}")
            raise ValueError(f"Not a Word file: {file_path}")

        logger.info(f"Loading Word document: {path.name}")

        try:
            doc = Document(str(path))
            documents = []
            full_text = []

            # ── Extract paragraphs ───────────────────────────────
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # ── Extract tables ───────────────────────────────────
            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_text.append(" | ".join(row_data))

                if table_text:
                    full_text.append(
                        f"\n[Table {table_num + 1}]\n" +
                        "\n".join(table_text)
                    )

            # ── Combine into one document ────────────────────────
            combined_text = "\n".join(full_text).strip()

            if combined_text:
                documents.append({
                    "content": combined_text,
                    "metadata": {
                        "source":    path.name,
                        "file_path": str(path),
                        "file_type": "docx",
                        "title":     path.stem,
                        "author":    "Unknown",
                    }
                })

            logger.info(f"Successfully extracted content from {path.name}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load Word file {path.name}: {str(e)}")
            raise

    def load_from_bytes(self, file_bytes: bytes,
                        filename: str) -> List[Dict[str, Any]]:
        """
        Load Word document from uploaded bytes (Streamlit uploads).

        Args:
            file_bytes: Raw bytes of the Word file
            filename:   Original filename
        """
        import io
        logger.info(f"Loading Word document from bytes: {filename}")

        try:
            doc = Document(io.BytesIO(file_bytes))
            documents = []
            full_text = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_data = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if row_data:
                        table_text.append(" | ".join(row_data))
                if table_text:
                    full_text.append(
                        f"\n[Table {table_num + 1}]\n" +
                        "\n".join(table_text)
                    )

            combined_text = "\n".join(full_text).strip()

            if combined_text:
                documents.append({
                    "content": combined_text,
                    "metadata": {
                        "source":    filename,
                        "file_type": "docx",
                        "title":     filename,
                        "author":    "Unknown",
                    }
                })

            logger.info(f"Extracted content from {filename}")
            return documents

        except Exception as e:
            logger.error(f"Failed to load Word bytes {filename}: {str(e)}")
            raise